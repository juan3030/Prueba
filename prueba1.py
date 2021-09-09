import copy
import os
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Cm

DIR = 'C:/Users/palac/Desktop/Proyecto Backend'
contenido = os.listdir(DIR)
imagenes = []
for fichero in contenido:
    if os.path.isfile(os.path.join(DIR, fichero)) and fichero.endswith('.jpg'):
        imagenes.append(fichero)
print(imagenes)

document = Document()
sections = document.sections
section = sections[0]

for i in imagenes:
    img = Image.open(i)
    img2 = copy.deepcopy(img)

    TamañoOriginal = img.size
    print (TamañoOriginal)

    size_max_small = (1123, 796)
    size_max_small2 = (796, 1123)

    img.thumbnail(size_max_small)
    Tamaño1 = img.size
    print (Tamaño1)

    img2.thumbnail(size_max_small2)
    Tamaño2 = img2.size
    print (Tamaño2)

    section.page_width = 7560000
    section.page_height = 10692000

    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)

    if Tamaño2 > Tamaño1:
        print ("Se pegara la imagen de manera Vertical")
        img2.save("Modificada"+i)
        document.add_picture(i, width=Cm(((Tamaño2[0])*2.54) / 96),
        height=Cm(((Tamaño2[1])*2.54) / 96))
    else:
        print ("Se pegara la imagen de manera horizontal")
        img.save("Modificada"+i)

        if Tamaño1[0] > Tamaño1[1]:
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
        document.add_picture(i, width=Cm(((Tamaño1[0])*2.54) / 96),
        height=Cm(((Tamaño1[1])*2.54) / 96))

    document.save('Final.docx')
    current_section = document.sections[-1]
    #current_section.start_type
    new_section = document.add_section(WD_SECTION.ODD_PAGE)
    #new_section.start_type

    Resultado = Image.open(i)
    print (Resultado.size)
